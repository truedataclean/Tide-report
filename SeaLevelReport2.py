import csv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import calendar
from datetime import datetime, timedelta
from docx2pdf import convert
import os

def read_csv(file_path):
    # Open the CSV file and read its content
    data = []
    with open(file_path, mode='r') as file:
        csv_reader = csv.reader(file)
        # Read the first line as file info
        file_info = next(csv_reader)
        # Skip the next 2 lines
        for _ in range(2):
            next(csv_reader)
        # Manually specify the header
        header = ["date", "day", "month", "year", "t1", "d1", "t2", "d2", "t3", "d3", "t4", "d4", "t5", "d5"]
        for row in csv_reader:
            data.append(row)
    return file_info, header, data

def is_daylight_saving(date):
    """
    Check if a given date falls within New Zealand's daylight saving time.
    Daylight saving in New Zealand starts at 2:00 AM on the last Sunday in September
    and ends at 3:00 AM on the first Sunday in April.
    """

    year = date.year

    # Calculate the last Sunday in September
    september_last_sunday = max(
        datetime(year, 9, day) for day in range(22, 30 + 1) if datetime(year, 9, day).weekday() == 6
    )

    # Calculate the first Sunday in April
    april_first_sunday = min(
        datetime(year, 4, day) for day in range(1, 7 + 1) if datetime(year, 4, day).weekday() == 6
    )

    # Check if the date is within the daylight saving period
    return september_last_sunday <= date < april_first_sunday

def find_new_zealand_daylight_saving_time(year):
    """
    Find the start and end dates of New Zealand's daylight saving time for a given year.
    Daylight saving in New Zealand starts at 2:00 AM on the last Sunday in September
    and ends at 3:00 AM on the first Sunday in April.
    """
    # Calculate the last Sunday in September
    september_last_sunday = max(
        datetime(year, 9, day) for day in range(22, 30 + 1) if datetime(year, 9, day).weekday() == 6
    )

    # Calculate the first Sunday in April
    april_first_sunday = min(
        datetime(year, 4, day) for day in range(1, 7 + 1) if datetime(year, 4, day).weekday() == 6
    )

    return september_last_sunday, april_first_sunday

def group_data_by_month(data):
    """Group data by month."""
    grouped_data = {}
    for row in data:
        month = row[2]  # Assuming the third column is the month
        if month not in grouped_data:
            grouped_data[month] = []
        grouped_data[month].append(row)
    return grouped_data

# add the top table
def add_top_table(doc):
    top_table = doc.add_table(rows=1, cols=2)
        # top_table.style = 'Table Grid'

        # Add logo to the first cell
    top_table.cell(0, 0).paragraphs[0].add_run().add_picture(r'N:\Publications\Toitū Te Whenua LINZ logo\toitu_te_whenua_colour_cmyk_66mm_png.png', width=Pt(180))
    top_table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(0)  # Remove blank space after paragraph
    top_table.cell(0, 0).vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER 

        # Add text and hyperlinks to the second cell
    # First paragraph
    text_paragraph1 = top_table.cell(0, 1).paragraphs[0]
    text_content1 = [
        ("Sourced from ", None),
        ("http://www.linz.govt.nz", RGBColor(0, 0, 255))
    ]
    for text, color in text_content1:
        run = text_paragraph1.add_run(text)
        run.font.size = Pt(10)
        if color:
            run.font.color.rgb = color
            run.font.underline = True
    text_paragraph1.paragraph_format.space_after = Pt(5)  # Reduce spacing after the paragraph
    text_paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align text to the left

    # Second paragraph
    text_paragraph2 = top_table.cell(0, 1).add_paragraph()
    text_content2 = [
        ("E-mail address ", None),
        ("hydro@linz.govt.nz", RGBColor(0, 0, 255))
    ]
    for text, color in text_content2:
        run = text_paragraph2.add_run(text)
        run.font.size = Pt(10)
        if color:
            run.font.color.rgb = color
            run.font.underline = True
    text_paragraph2.paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
    text_paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align text to the left

    # Set vertical alignment for the cell
    top_table.cell(0, 1).vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align cell content vertically to the middle
    

# Function to add a title line to the document
def add_title(doc):
    title_paragraph = doc.add_paragraph("New Zealand Hydrographic Authority Tide Predictions")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0]
    title_run.font.size = Pt(10)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    title_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title

def add_title1(doc):
    title_paragraph = doc.add_paragraph("New Zealand Hydrographic Authority Tide Stream Predictions")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0]
    title_run.font.size = Pt(10)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    title_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title

def add_header(doc, regionname):
    header_paragraph = doc.add_paragraph(regionname)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_paragraph.runs[0]
    header_run.font.size = Pt(20)
    header_run.font.bold = True
    header_run.font.name = 'Arial'
    header_run.font.color.rgb = RGBColor(20, 171, 155)  # Set color to #14ab9b
    header_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the header

def add_coordinates(doc, coord):
    ldt_paragraph = doc.add_paragraph(coord)
    ldt_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ldt_run = ldt_paragraph.runs[0]
    ldt_run.font.size = Pt(10)
    ldt_run.font.name = 'Arial'
    ldt_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the header

def add_month_heading(document, month, rows):
    """Add a month and year heading to the document."""
    month_name = calendar.month_name[int(month)]
    # Add year heading from the first data row of the month
    year = rows[0][3]

    heading = document.add_heading(f"{month_name} {year}", level=1)
    for run in heading.runs:
        run.font.size = Pt(20)  # Set font size to 24
        run.font.color.rgb = RGBColor(20, 171, 155)  # Set font color to RGB(20, 171, 155)
        run.font.bold = True
        run.font.name = 'Arial'

    heading.paragraph_format.space_before = Pt(0)  # Reduce spacing under the heading
    heading.paragraph_format.space_after = Pt(0)  # Reduce spacing under the heading
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Function to add a condition line to the document
def add_condition(doc):
    condition_paragraph = doc.add_paragraph()
    condition_run = condition_paragraph.add_run("Tidal Stream ")
    condition_run.font.size = Pt(10)
    condition_run.font.name = 'Arial'

    begins_run = condition_paragraph.add_run("begins")
    begins_run.font.underline = True
    begins_run.font.size = Pt(10)
    begins_run.font.name = 'Arial'

    condition_run = condition_paragraph.add_run(" at the N.Z. Local Time shown, in the direction indicated")
    condition_run.font.size = Pt(10)
    condition_run.font.name = 'Arial'

    condition_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    condition_paragraph.paragraph_format.space_after = Pt(5)  # Reduce spacing after the title

def add_condition1(doc):
    condition_paragraph = doc.add_paragraph("Chatham Islands Local Times and Heights of High and Low Waters")
    condition_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    condition_run = condition_paragraph.runs[0]
    condition_run.font.size = Pt(10)
    condition_run.font.name = 'Arial'
    condition_paragraph.paragraph_format.space_after = Pt(5)  # Reduce spacing after the title

def add_condition2(doc):
    condition_paragraph = doc.add_paragraph("N.Z. Local Times and Heights of High and Low Waters")
    condition_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    condition_run = condition_paragraph.runs[0]
    condition_run.font.size = Pt(10)
    condition_run.font.name = 'Arial'
    condition_paragraph.paragraph_format.space_after = Pt(5)  # Reduce spacing after the title

def add_caution(doc):
    caution_paragraph = doc.add_paragraph("Caution: Tidal Streams may be subject to irregularities and these times should be regarded as approximate only.")
    caution_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caution_run = caution_paragraph.runs[0]
    caution_run.font.size = Pt(8.5)
    caution_run.font.name = 'Arial'
    caution_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title
    caution_paragraph.paragraph_format.space_before = Pt(5)  # Reduce spacing before the title

def add_daylight(doc):
    daylight_paragraph = doc.add_paragraph("Times listed are N.Z. Daylight Time")
    daylight_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    daylight_run = daylight_paragraph.runs[0]
    daylight_run.font.size = Pt(10)
    daylight_run.font.name = 'Arial'
    daylight_paragraph.paragraph_format.space_after = Pt(3)  # Reduce spacing after the title

def add_daylight1(doc):
    daylight_paragraph = doc.add_paragraph("Times listed are Chatham Islands Standard Time")
    daylight_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    daylight_run = daylight_paragraph.runs[0]
    daylight_run.font.size = Pt(10)
    daylight_run.font.name = 'Arial'
    daylight_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title

def add_copyright(doc):
    copyright_paragraph = doc.add_paragraph("Crown Copyright Reserved")
    copyright_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_run = copyright_paragraph.runs[0]
    copyright_run.font.size = Pt(10)
    copyright_run.font.name = 'Arial'
    copyright_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the title


def save_to_word(file_info, grouped_data, output_path):
    # Add region name and coordinates
    region_name = file_info[1]
    coordinates = f"Lat. {file_info[2]} Long. {file_info[3]}"
    """Save grouped data to a Word document."""
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'  # Set font to Arial
    font.size = Pt(10)  # Set font size to 12
    
    # Add grouped data, each month on a separate page
    first_page = True  # Flag to track the first page
    for month, rows in grouped_data.items():
        if not first_page:
            document.add_page_break()  # Add a page break for each month
        first_page = False  # Set the flag to False after the first page

        add_top_table(document)
        if region_name == "Te Aumiti / French Pass" or region_name == "Tory Channel / Kura Te Au Entrance":
            add_title1(document)
        else:
            add_title(document)
        add_header(document, region_name)
        add_coordinates(document, coordinates)
        add_month_heading(document, month, rows)
        if region_name == "Owenga - Chatham Island" or region_name == "Kaingaroa - Chatham Island" or region_name == "Waitangi - Chatham Island":
            add_condition1(document)
        elif region_name == "Te Aumiti / French Pass" or region_name == "Tory Channel / Kura Te Au Entrance":
            add_condition(document)
        else:
            add_condition2(document)

        # Add data table
        table = document.add_table(rows=9, cols=12)
        # table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        if region_name == "Te Aumiti / French Pass" or region_name == "Tory Channel / Kura Te Au Entrance":
            headers = ['', 'Time', 'Dir', '', 'Time', 'Dir', '', 'Time', 'Dir', '', 'Time', 'Dir']
        else:
            headers = ['', 'Time', 'm', '', 'Time', 'm', '', 'Time', 'm', '', 'Time', 'm']

        # Set column widths
        # for col in table.columns:
        #     for cell in col.cells:
        #         cell.width = Pt(50)

        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing under the header
                paragraph.paragraph_format.space_before = Pt(0)  # Reduce spacing above the header
                paragraph.paragraph_format._line_spacing_rule = 0  # Set line spacing to single
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Set header font size to 10
        
        
        # Add data rows
        for row in rows:
            # print(row)  # Print the row for debuggin

            # Extract date, day, and rest of the data from the row
            date = row[0]  # Assuming the first column is the date
            day = row[1]   # Assuming the second column is the day
            rest = row[4:] # The rest of the columns
            year = int(rows[0][3])

            if int(month) == 4 and int(date) == 1:
                start_dst, end_dst = find_new_zealand_daylight_saving_time(year)  # Call the function to find daylight saving time
                # print(f"Daylight Saving Time starts ends on: {end_dst}")  # Print the start and end dates
                # print(f"Year: {year}, Month: {month}, Day: {day}")  # Print the year, month, and day for debugging
                end_dst_day = end_dst.day  # Extract the day number from the datetime object
                # print(f"End DST Day: {end_dst_day}")  # Print the extracted day number

            elif int(month) == 9 and int(date) == 1:
                start_dst, end_dst = find_new_zealand_daylight_saving_time(year)  # Call the function to find daylight saving time
                # print(f"Daylight Saving Time starts on: {start_dst}")  # Print the start and end dates
                # print(f"Year: {year}, Month: {month}, Day: {day}")  # Print the year, month, and day for debugging
                start_dst_day = start_dst.day  # Extract the day number from the datetime object
                # print(f"Start DST Day: {start_dst_day}")  # Print the extracted day number

            for idx, row_date in enumerate(range(1, 9)):  # Changed range to 1 to 9 for 8 days
                if date == str(row_date):
                    # Ensure the table has enough rows
                    while len(table.rows) <= idx + 1:
                        table.add_row()
                    # Fill the first column with the first 8 days
                    target_row = table.rows[idx + 1].cells
                    target_row[0].text = f"{date}"
                    run_date = target_row[0].paragraphs[0].runs[0]
                    run_date.font.size = Pt(22)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[0].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(11)
                    target_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[0].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[0].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

                    target_row[1].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[1].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[1].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[1].paragraphs[0].paragraph_format.space_after = Pt(0)

                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[1].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                        target_row[1].paragraphs[0].paragraph_format.space_after = Pt(0)

                    elif int(month) == 4 and int(date) < end_dst_day:
                        # print(f"Date: {date}, Month: {month}, Day: {day}", {end_dst_day})  # Print the year, month, and day for debugging
                        target_row[1].paragraphs[0].runs[0].font.bold = True  # Make font bold for April
                        target_row[1].paragraphs[0].paragraph_format.space_after = Pt(0)

                    elif int(month) == 4 and int(date) == end_dst_day:
                        target_row[1].paragraphs[0].runs[0].font.bold = True
                        # Calculate rows where time is before 3 AM
                        times = [rest[i] for i in range(0, len(rest), 2)]
                        before_3am_count = sum(1 for time in times if time and int(time.split(':')[0]) < 3)
                        # Update the text for the cell based on the row count
                        target_row[1].paragraphs[0].runs[0].text = "\n".join(
                            [rest[i] for i in range(0, len(rest), 2)][:before_3am_count]
                        )
                        target_row[1].paragraphs[0].paragraph_format.space_after = Pt(0)
                        # Add the rest of the range with non-bold style
                        remaining_times = [rest[i] for i in range(0, len(rest), 2)][before_3am_count:]
                        if remaining_times:
                            non_bold_paragraph = target_row[1].add_paragraph("\n".join(remaining_times))
                            non_bold_paragraph.runs[0].font.size = Pt(9)
                            non_bold_paragraph.paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                            non_bold_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
        

                    target_row[2].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[2].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[2].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[2].paragraphs[0].paragraph_format.space_after = Pt(8)  # Reduce spacing after the paragraph

                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[2].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months

                    elif int(month) == 4 and int(date) < end_dst_day:
                        # print(f"Date: {date}, Month: {month}, Day: {day}", {end_dst_day})  # Print the year, month, and day for debugging
                        target_row[2].paragraphs[0].runs[0].font.bold = True  # Make font bold for April

                    elif int(month) == 4 and int(date) == end_dst_day:
                        target_row[2].paragraphs[0].runs[0].font.bold = True

                        # Calculate rows where time is before 3 AM
                        times = [rest[i] for i in range(0, len(rest), 2)]
                        before_3am_count = sum(1 for time in times if time and int(time.split(':')[0]) < 3)
                        # Update the text for the cell based on the row count
                        target_row[2].paragraphs[0].runs[0].text = "\n".join(
                            [rest[i] for i in range(1, len(rest), 2)][:before_3am_count]
                        )
                        # Add the rest of the range with non-bold style
                        remaining_times = [rest[i] for i in range(1, len(rest), 2)][before_3am_count:]
                        if remaining_times:
                            non_bold_paragraph = target_row[2].add_paragraph("\n".join(remaining_times))
                            non_bold_paragraph.runs[0].font.size = Pt(9)
                            non_bold_paragraph.paragraph_format.line_spacing_rule = 0  # Set line spacing to single   
                            non_bold_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    

            for idx, row_date in enumerate(range(9, 17)):  # Changed range to 9 to 16 for the next 8 days
                if date == str(row_date):
                    # Fill the second column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the second column
                    target_row[3].text = f"{date}"
                    run_date = target_row[3].paragraphs[0].runs[0]
                    run_date.font.size = Pt(22)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[3].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(11)
                    target_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[3].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[3].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    target_row[4].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[4].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[4].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[4].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[4].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                    
                    target_row[5].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[5].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[5].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[5].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[5].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months

            for idx, row_date in enumerate(range(17, 25)):  # Changed range to 17 to 24 for the next 8 days
                if date == str(row_date):
                    # Fill the third column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the third column
                    target_row[6].text = f"{date}"
                    run_date = target_row[6].paragraphs[0].runs[0]
                    run_date.font.size = Pt(22)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[6].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(11)
                    target_row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[6].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[6].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    target_row[7].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[7].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[7].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[7].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[7].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                    
                    target_row[8].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[8].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[8].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[8].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[8].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months

            for idx, row_date in enumerate(range(25, 32)):  # Changed range to 25 to 31 for the next 8 days
                if date == str(row_date):
                    # Fill the fourth column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the fourth column
                    target_row[9].text = f"{date}"
                    run_date = target_row[9].paragraphs[0].runs[0]
                    run_date.font.size = Pt(22)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[9].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(11)
                    target_row[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[9].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[9].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

                    target_row[10].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[10].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[10].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[10].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[10].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                    elif int(month) == 9 and int(date) > start_dst_day:
                        # print(f"Date: {date}, Month: {month}, Day: {day}", {start_dst_day})
                        target_row[10].paragraphs[0].runs[0].font.bold = True  # Make font bold for September
                    elif int(month) == 9 and int(date) == start_dst_day:
        
                        # Calculate rows where time is after 2 AM
                        times = [rest[i] for i in range(0, len(rest), 2)]
                        before_2am_count = sum(1 for time in times if time and int(time.split(':')[0]) < 2)
                        # Update the text for the cell based on the row count
                        target_row[10].paragraphs[0].runs[0].text = "\n".join(
                            [rest[i] for i in range(0, len(rest), 2)][:before_2am_count]
                        )
                        # Add the rest of the range with non-bold style
                        remaining_times = [rest[i] for i in range(0, len(rest), 2)][before_2am_count:]
                        if remaining_times:
                            non_bold_paragraph = target_row[10].add_paragraph("\n".join(remaining_times))
                            non_bold_paragraph.runs[0].font.size = Pt(9)
                            non_bold_paragraph.runs[0].font.bold = True  # Make font bold
                            non_bold_paragraph.paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                            non_bold_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

                    target_row[11].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[11].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[11].paragraphs[0].paragraph_format.line_spacing_rule = 0
                    target_row[11].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[11].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                    elif int(month) == 9 and int(date) > start_dst_day:
                        # print(f"Date: {date}, Month: {month}, Day: {day}", {start_dst_day})
                        target_row[11].paragraphs[0].runs[0].font.bold = True  # Make font bold for September

                    elif int(month) == 9 and int(date) == start_dst_day:
        
                        # Calculate rows where time is after 2 AM
                        times = [rest[i] for i in range(0, len(rest), 2)]
                        before_2am_count = sum(1 for time in times if time and int(time.split(':')[0]) < 2)
                        # Update the text for the cell based on the row count
                        target_row[11].paragraphs[0].runs[0].text = "\n".join(
                            [rest[i] for i in range(1, len(rest), 2)][:before_2am_count]
                        )
                        # Add the rest of the range with non-bold style
                        remaining_times = [rest[i] for i in range(1, len(rest), 2)][before_2am_count:]
                        if remaining_times:
                            non_bold_paragraph = target_row[11].add_paragraph("\n".join(remaining_times))
                            non_bold_paragraph.runs[0].font.size = Pt(9)
                            non_bold_paragraph.runs[0].font.bold = True  # Make font bold
                            non_bold_paragraph.paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                            non_bold_paragraph.paragraph_format.space_after = Pt(0)


        add_caution(document)  # Add caution line after the table 
        if region_name == "Owenga - Chatham Island" or region_name == "Kaingaroa - Chatham Island" or region_name == "Waitangi - Chatham Island":
            add_daylight1(document)  # Add Chatham Islands daylight line after the table
        else:
            add_daylight(document)  # Add NZ daylight line after the table

        add_copyright(document)  # Add copyright line after the table
    
    # Save the document
    document.save(output_path)

def convert_to_pdf(docx_path, pdf_path):
        """
        Convert a Word document to a PDF file.
        """
        convert(docx_path, pdf_path)

def main():
    """Main function to execute the script."""
    # Define the folder path containing CSV files
    folder_path = r'C:\Projects\Glen\CSV files'
    output_folder = r'C:\Projects\Glen\Reports'
    os.makedirs(output_folder, exist_ok=True)

    # Process each CSV file in the folder
    for file in os.listdir(folder_path):
        if file.endswith('.csv'):
            file_path = os.path.join(folder_path, file)
            output_path = os.path.join(output_folder, os.path.splitext(file)[0] + '.docx')
            pdf_path = os.path.join(output_folder, os.path.splitext(file)[0] + '.pdf')

            # Read the CSV file
            file_info, header, data = read_csv(file_path)

            # Group data by month
            grouped_data = group_data_by_month(data)

            # Save grouped data to a Word document
            save_to_word(file_info, grouped_data, output_path)

            # Convert the Word document to PDF
            convert_to_pdf(output_path, pdf_path)

            print(f"Processed: {file}")
            print(f"Word document saved to {output_path}")
            print(f"PDF document saved to {pdf_path}")

if __name__ == "__main__":
    main()