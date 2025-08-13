"""Main module."""
import csv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import calendar
from datetime import datetime, timedelta
from docx2pdf import convert
import os
import yaml
from docx.shared import Mm
from tkinter import Tk, simpledialog
from tkinter import messagebox
import unicodedata

def read_csv(file_path):
    """
    Reads a CSV file and extracts file information, header, and data.

    Args:
        file_path (str): Path to the CSV file.

    Returns:
        tuple: A tuple containing file_info (list), header (list), and data (list of lists).

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file format is invalid or missing required data.
    """
    try:
        with open(file_path, mode='r') as file:
            csv_reader = csv.reader(file)
            try:
                file_info = next(csv_reader)  # Read the first line as file info
            except StopIteration:
                raise ValueError("CSV file is empty or missing file information.")

            # Skip the next 2 lines
            for _ in range(2):
                try:
                    next(csv_reader, None)
                except StopIteration:
                    raise ValueError("CSV file is missing required header or data rows.")

            # Define the header manually
            header = ["date", "day", "month", "year", "t1", "d1", "t2", "d2", "t3", "d3", "t4", "d4", "t5", "d5"]

            # Read the remaining rows as data
            data = list(csv_reader)
            if not data:
                raise ValueError("CSV file contains no data rows.")

        return file_info, header, data

    except FileNotFoundError:
        raise FileNotFoundError(f"The file at path '{file_path}' does not exist.")
    except Exception as e:
        raise ValueError(f"An error occurred while reading the CSV file: {e}")

def is_daylight_saving(date):
    """
    Check if a given date falls within New Zealand's daylight saving time.
    Daylight saving in New Zealand starts at 2:00 AM on the last Sunday in September
    and ends at 3:00 AM on the first Sunday in April.

    Args:
        date (datetime): The date to check.

    Returns:
        bool: True if the date is within daylight saving time, False otherwise.

    Raises:
        ValueError: If the input is not a datetime object.
    """
    if not isinstance(date, datetime):
        raise ValueError("The input must be a datetime object.")

    try:
        year = date.year

        # Calculate the last Sunday in September
        september_last_sunday = max(
            datetime(year, 9, day) for day in range(22, 30 + 1) if datetime(year, 9, day).weekday() == 6
        )

        # Calculate the first Sunday in April
        april_first_sunday = min(
            datetime(year, 4, day) for day in range(1, 7 + 1) if datetime(year, 4, day).weekday() == 6
        )

        # Check if the date is within the daylight saving period
        return september_last_sunday <= date < april_first_sunday

    except Exception as e:
        raise ValueError(f"An error occurred while determining daylight saving time: {e}")

def find_new_zealand_daylight_saving_time(year):
    """
    Find the start and end dates of New Zealand's daylight saving time for a given year.
    Daylight saving in New Zealand starts at 2:00 AM on the last Sunday in September
    and ends at 3:00 AM on the first Sunday in April.

    Args:
        year (int): The year for which to calculate daylight saving time.

    Returns:
        tuple: A tuple containing the start and end dates of daylight saving time.

    Raises:
        ValueError: If the input year is not a valid integer or is out of range.
    """
    if not isinstance(year, int) or year < 1:
        raise ValueError("The year must be a positive integer.")

    try:
        # Calculate the last Sunday in September
        september_last_sunday = max(
            datetime(year, 9, day) for day in range(22, 30 + 1) if datetime(year, 9, day).weekday() == 6
        )

        # Calculate the first Sunday in April
        april_first_sunday = min(
            datetime(year, 4, day) for day in range(1, 7 + 1) if datetime(year, 4, day).weekday() == 6
        )

        return september_last_sunday, april_first_sunday

    except Exception as e:
        raise ValueError(f"An error occurred while calculating daylight saving time: {e}")

def group_data_by_month(data):
    """
    Group data by month.

    Args:
        data (list): List of data rows.

    Returns:
        dict: A dictionary where keys are months and values are lists of rows for that month.

    Raises:
        ValueError: If the data format is invalid or missing required columns.
    """
    if not isinstance(data, list):
        raise ValueError("Input data must be a list of rows.")

    grouped_data = {}
    try:
        for row in data:
            if len(row) < 3:
                raise ValueError("Each row must have at least 3 columns (date, day, month).")
            month = row[2]  # Assuming the third column is the month
            if not month.isdigit() or not (1 <= int(month) <= 12):
                raise ValueError(f"Invalid month value: {month}")
            if month not in grouped_data:
                grouped_data[month] = []
            grouped_data[month].append(row)
    except Exception as e:
        raise ValueError(f"An error occurred while grouping data by month: {e}")

    return grouped_data

# add the top table
def add_top_table(doc, linz_logo_path):
    # Add LINZ logo and text to the top of the document - the system cannot process Māori characters correctly
    # Therefore, the image path is hardcoded
    try:
        if not linz_logo_path or any(char in linz_logo_path for char in ['ā', 'ē', 'ī', 'ō', 'Ū', 'Ā', 'Ē', 'Ō']):
            linz_logo_path = 'N:\\Publications\\Toitū Te Whenua LINZ logo\\toitu_te_whenua_colour_cmyk_66mm_png.png'
        else:
            linz_logo_path = linz_logo_path

        if not os.path.exists(linz_logo_path):
            raise FileNotFoundError(f"LINZ logo file not found at path: {linz_logo_path}")
    except Exception as e:
        raise ValueError(f"An error occurred while validating the LINZ logo path: {e}")
    try:
        top_table = doc.add_table(rows=1, cols=2)
        # Add logo to the first cell
        try:
            if os.path.exists(linz_logo_path):
                top_table.cell(0, 0).paragraphs[0].add_run().add_picture(linz_logo_path, width=Pt(180))
            else:
                raise FileNotFoundError(f"Logo file not found at path: {linz_logo_path}")
        except FileNotFoundError:
            raise FileNotFoundError("Logo file not found. Please check the file path.")
        except Exception as e:
            raise ValueError(f"An error occurred while adding the logo: {e}")

        top_table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(0)  # Remove blank space after paragraph
        top_table.cell(0, 0).vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add text and hyperlinks to the second cell
        try:
            # First paragraph
            text_paragraph1 = top_table.cell(0, 1).paragraphs[0]
            text_content1 = [
                ("Sourced from ", None),
                ("http://www.linz.govt.nz", RGBColor(0, 0, 255))
            ]
            for text, color in text_content1:
                run = text_paragraph1.add_run(text)
                run.font.size = Pt(10)
                if color:
                    run.font.color.rgb = color
                    run.font.underline = True
            text_paragraph1.paragraph_format.space_after = Pt(5)  # Reduce spacing after the paragraph
            text_paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align text to the left

            # Second paragraph
            text_paragraph2 = top_table.cell(0, 1).add_paragraph()
            text_content2 = [
                ("E-mail address ", None),
                ("hydro@linz.govt.nz", RGBColor(0, 0, 255))
            ]
            for text, color in text_content2:
                run = text_paragraph2.add_run(text)
                run.font.size = Pt(10)
                if color:
                    run.font.color.rgb = color
                    run.font.underline = True
            text_paragraph2.paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
            text_paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align text to the left

            # Set vertical alignment for the cell
            top_table.cell(0, 1).vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align cell content vertically to the middle
        except Exception as e:
            raise ValueError(f"An error occurred while adding text and hyperlinks: {e}")
    except Exception as e:
        raise ValueError(f"An error occurred while creating the top table: {e}")

# Function to add a title line to the document
def add_title(doc):
    try:
        title_paragraph = doc.add_paragraph("New Zealand Hydrographic Authority Tide Predictions")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(10)
        title_run.font.bold = True
        title_run.font.name = 'Arial'
        title_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the title: {e}")

def add_title1(doc):
    try:
        title_paragraph = doc.add_paragraph("New Zealand Hydrographic Authority Tide Stream Predictions")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(10)
        title_run.font.bold = True
        title_run.font.name = 'Arial'
        title_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the title: {e}")

def add_header(doc, regionname):
    try:
        header_paragraph = doc.add_paragraph(regionname)
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_paragraph.runs[0]
        header_run.font.size = Pt(20)
        header_run.font.bold = True
        header_run.font.name = 'Arial'
        header_run.font.color.rgb = RGBColor(20, 171, 155)  # Set color to #14ab9b
        header_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the header
    except Exception as e:
        raise ValueError(f"An error occurred while adding the header: {e}")

def add_coordinates(doc, coord):
    try:
        ldt_paragraph = doc.add_paragraph(coord)
        ldt_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ldt_run = ldt_paragraph.runs[0]
        ldt_run.font.size = Pt(10)
        ldt_run.font.name = 'Arial'
        ldt_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the header
    except Exception as e:
        raise ValueError(f"An error occurred while adding coordinates: {e}")

def add_month_heading(document, month, rows):
    """
    Add a month and year heading to the document.

    Args:
        document (Document): The Word document object.
        month (str): The month number as a string.
        rows (list): List of rows containing data for the month.

    Raises:
        ValueError: If the month is invalid or rows are missing required data.
    """
    try:
        if not month.isdigit() or not (1 <= int(month) <= 12):
            raise ValueError(f"Invalid month value: {month}")

        if not rows or len(rows[0]) < 4:
            raise ValueError("Rows are missing required data for year extraction.")

        month_name = calendar.month_name[int(month)]
        year = rows[0][3]  # Extract year from the first data row

        heading = document.add_heading(f"{month_name} {year}", level=1)
        for run in heading.runs:
            run.font.size = Pt(20)  # Set font size to 24
            run.font.color.rgb = RGBColor(20, 171, 155)  # Set font color to RGB(20, 171, 155)
            run.font.bold = True
            run.font.name = 'Arial'

        heading.paragraph_format.space_before = Pt(0)  # Reduce spacing under the heading
        heading.paragraph_format.space_after = Pt(0)  # Reduce spacing under the heading
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception as e:
        raise ValueError(f"An error occurred while adding the month heading: {e}")


# Function to add a condition line to the document
def add_condition(doc):
    try:
        condition_paragraph = doc.add_paragraph()
        condition_run = condition_paragraph.add_run("Tidal Stream ")
        condition_run.font.size = Pt(10)
        condition_run.font.name = 'Arial'

        begins_run = condition_paragraph.add_run("begins")
        begins_run.font.underline = True
        begins_run.font.size = Pt(10)
        begins_run.font.name = 'Arial'

        condition_run = condition_paragraph.add_run(" at the N.Z. Local Time shown, in the direction indicated")
        condition_run.font.size = Pt(10)
        condition_run.font.name = 'Arial'

        condition_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        condition_paragraph.paragraph_format.space_after = Pt(5)  # Reduce spacing after the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the condition: {e}")

def add_condition1(doc):
    try:
        condition_paragraph = doc.add_paragraph("Chatham Islands Local Times and Heights of High and Low Waters")
        condition_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        condition_run = condition_paragraph.runs[0]
        condition_run.font.size = Pt(10)
        condition_run.font.name = 'Arial'
        condition_paragraph.paragraph_format.space_after = Pt(5)  # Reduce spacing after the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the condition: {e}")

def add_condition2(doc):
    try:
        condition_paragraph = doc.add_paragraph("N.Z. Local Times and Heights of High and Low Waters")
        condition_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        condition_run = condition_paragraph.runs[0]
        condition_run.font.size = Pt(10)
        condition_run.font.name = 'Arial'
        condition_paragraph.paragraph_format.space_after = Pt(5)  # Reduce spacing after the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the condition: {e}")

def add_caution(doc):
    try:
        caution_paragraph = doc.add_paragraph("Caution: Tidal Streams may be subject to irregularities and these times should be regarded as approximate only.")
        caution_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caution_run = caution_paragraph.runs[0]
        caution_run.font.size = Pt(8.5)
        caution_run.font.name = 'Arial'
        caution_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title
        caution_paragraph.paragraph_format.space_before = Pt(5)  # Reduce spacing before the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the caution: {e}")

def add_daylight(doc, month):
    try:
        if int(month) == 4 or int(month) == 9:
            daylight_paragraph = doc.add_paragraph("Times shown in bold have been adjusted for N.Z. Daylight Time")
        elif int(month) in [1, 2, 3, 10, 11, 12]:
            daylight_paragraph = doc.add_paragraph("Times listed are N.Z. Daylight Time")
        else:
            daylight_paragraph = doc.add_paragraph("Times listed are N.Z. Standard Time")
        daylight_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        daylight_run = daylight_paragraph.runs[0]
        daylight_run.font.size = Pt(10)
        daylight_run.font.name = 'Arial'
        daylight_paragraph.paragraph_format.space_after = Pt(3)  # Reduce spacing after the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the daylight paragraph: {e}")

def add_daylight1(doc, month):
    try:
        if int(month) == 4 or int(month) == 9:
            daylight_paragraph = doc.add_paragraph("Times shown in bold have been adjusted for Chatham Islands Daylight Time")
        elif int(month) in [1, 2, 3, 10, 11, 12]:
            daylight_paragraph = doc.add_paragraph("Times listed are Chatham Islands Daylight Time")
        else:
            daylight_paragraph = doc.add_paragraph("Times listed are Chatham Islands Standard Time")
        daylight_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        daylight_run = daylight_paragraph.runs[0]
        daylight_run.font.size = Pt(10)
        daylight_run.font.name = 'Arial'
        daylight_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the Chatham Islands daylight paragraph: {e}")

def add_copyright(doc):
    try:
        copyright_paragraph = doc.add_paragraph("Crown Copyright Reserved")
        copyright_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        copyright_run = copyright_paragraph.runs[0]
        copyright_run.font.size = Pt(10)
        copyright_run.font.name = 'Arial'
        copyright_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the title
    except Exception as e:
        raise ValueError(f"An error occurred while adding the copyright paragraph: {e}")


def save_to_word(file_info, grouped_data, output_path, linz_logo_path):
    # Add region name and coordinates
    # Fix region name macrons as Māori
    # when UTF-8 encoded text is incorrectly decoded using single-byte encoding such as Latin-1 or Windows-1252."
    region_name = file_info[1]
    fixed = region_name.encode('Windows-1252').decode('utf-8')
    region_name = fixed
    Lat = (file_info[2]).replace('Â', '')
    Long = (file_info[3]).replace('Â', '')

    coordinates = f"Lat. {Lat} Long. {Long}"
    """Save grouped data to a Word document."""
    document = Document()
    style = document.styles['Normal']

    # Set document margins (in millimeters)
    sections = document.sections
    for section in sections:
        section.top_margin = Mm(15)      # 15 mm
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)

    font = style.font
    font.name = 'Arial'  # Set font to Arial
    font.size = Pt(10)  # Set font size to 12
    
    # Add grouped data, each month on a separate page
    first_page = True  # Flag to track the first page
    for month, rows in grouped_data.items():
        if not first_page:
            document.add_page_break()  # Add a page break for each month
        first_page = False  # Set the flag to False after the first page

        add_top_table(document, linz_logo_path)
        if region_name == "Te Aumiti / French Pass" or region_name == "Tory Channel / Kura Te Au Entrance":
            add_title1(document)
        else:
            add_title(document)
        add_header(document, region_name)
        add_coordinates(document, coordinates)
        add_month_heading(document, month, rows)
        if region_name == "Owenga - Chatham Island" or region_name == "Kaingaroa - Chatham Island" or region_name == "Waitangi - Chatham Island":
            add_condition1(document)
        elif region_name == "Te Aumiti / French Pass" or region_name == "Tory Channel / Kura Te Au Entrance":
            add_condition(document)
        else:
            add_condition2(document)
        
        # Add spacing before the table
        document.add_paragraph().paragraph_format.space_after = Pt(1)

        # Add data table
        table = document.add_table(rows=9, cols=12)
        # Center the table on the page
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        if region_name == "Te Aumiti / French Pass" or region_name == "Tory Channel / Kura Te Au Entrance":
            headers = ['', 'Time', 'Dir', '', 'Time', 'Dir', '', 'Time', 'Dir', '', 'Time', 'Dir']
        else:
            headers = ['', 'Time', 'm', '', 'Time', 'm', '', 'Time', 'm', '', 'Time', 'm']

        # Set column widths
        # for col in table.columns:
        #     for cell in col.cells:
        #         cell.width = Pt(50)

        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(5)  # Reduce spacing under the header
                paragraph.paragraph_format.space_before = Pt(0)  # Reduce spacing above the header
                paragraph.paragraph_format._line_spacing_rule = 0  # Set line spacing to single
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Set header font size to 10
        
        
        # Add data rows
        for row in rows:
            # print(row)  # Print the row for debuggin

            # Extract date, day, and rest of the data from the row
            date = row[0]  # Assuming the first column is the date
            day = row[1]   # Assuming the second column is the day
            rest = row[4:] # The rest of the columns
            year = int(rows[0][3])

            if int(month) == 4 and int(date) == 1:
                start_dst, end_dst = find_new_zealand_daylight_saving_time(year)  # Call the function to find daylight saving time
                # print(f"Daylight Saving Time starts ends on: {end_dst}")  # Print the start and end dates
                # print(f"Year: {year}, Month: {month}, Day: {day}")  # Print the year, month, and day for debugging
                end_dst_day = end_dst.day  # Extract the day number from the datetime object
                # print(f"End DST Day: {end_dst_day}")  # Print the extracted day number

            elif int(month) == 9 and int(date) == 1:
                start_dst, end_dst = find_new_zealand_daylight_saving_time(year)  # Call the function to find daylight saving time
                # print(f"Daylight Saving Time starts on: {start_dst}")  # Print the start and end dates
                # print(f"Year: {year}, Month: {month}, Day: {day}")  # Print the year, month, and day for debugging
                start_dst_day = start_dst.day  # Extract the day number from the datetime object
                # print(f"Start DST Day: {start_dst_day}")  # Print the extracted day number

            for idx, row_date in enumerate(range(1, 9)):  # Changed range to 1 to 9 for 8 days
                if date == str(row_date):
                    # Ensure the table has enough rows
                    while len(table.rows) <= idx + 1:
                        table.add_row()
                    # Fill the first column with the first 8 days
                    target_row = table.rows[idx + 1].cells
                    target_row[0].text = f"{date}"
                    run_date = target_row[0].paragraphs[0].runs[0]
                    run_date.font.size = Pt(22)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[0].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(11)
                    target_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[0].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[0].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

                    target_row[1].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[1].paragraphs[0].runs[0].font.size = Pt(10)
                    target_row[1].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[1].paragraphs[0].paragraph_format.space_after = Pt(0)

                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[1].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                        target_row[1].paragraphs[0].paragraph_format.space_after = Pt(0)

                    elif int(month) == 4 and int(date) < end_dst_day:
                        # print(f"Date: {date}, Month: {month}, Day: {day}", {end_dst_day})  # Print the year, month, and day for debugging
                        target_row[1].paragraphs[0].runs[0].font.bold = True  # Make font bold for April
                        target_row[1].paragraphs[0].paragraph_format.space_after = Pt(0)

                    elif int(month) == 4 and int(date) == end_dst_day:
                        target_row[1].paragraphs[0].runs[0].font.bold = False
                        bold_paragraph = target_row[1].paragraphs[0]
                        # Calculate rows where time is before 3 AM
                        times = [rest[i] for i in range(0, len(rest), 2)]
                        before_3am_count = sum(1 for time in times if time and int(time.split(':')[0]) < 3)
                        # Update the text for the cell based on the row count
                        ontimes = [rest[i] for i in range(0, len(rest), 2)][:before_3am_count] #range(0 to 1 for update the tide and direction)
                        remaining_times = [rest[i] for i in range(0, len(rest), 2)][before_3am_count:]


                        # Check if the first time is between 1:59 and 3:00 and prompt the user
                        no_condition = False  # Initialize no_condition variable

                        if times and times[0]:
                            try:
                                hour, minute = map(int, times[0].split(':'))
                                if (hour == 1 and minute >= 59) or (hour == 2) or (hour == 3 and minute == 0):
                                    root = Tk()
                                    root.withdraw()  # Hide the main window
                                    result = messagebox.askyesno(
                                        f"Confirmation Required {region_name}",
                                        f"{date}/{month}/{year} {times[0]}.\nIs this Daylight time?\n\nClick 'Yes' to NZDT, or 'No' to NZST.",
                                    )
                                    root.destroy()
                                    if result:
                                        # Yes condition: proceed as normal
                                        pass
                                    else:
                                        # No condition: skip or handle as needed
                                        no_condition = True
                                        # continue
                            except Exception as e:
                                print(f"Error parsing time '{times[0]}': {e}")

                        if ontimes:

                            bold_paragraph.clear()  # Clear existing text
                            bold_paragraph.runs.clear()  # Clear existing runs
                            run = bold_paragraph.add_run("\n".join(ontimes))
                            run.font.size = Pt(10)

                            if no_condition == True:
                                run.font.bold = False
                            else:
                                run.font.bold = True

                            for remaining_time in remaining_times:
                                if remaining_time:
                                    bold_paragraph.add_run(f"\n{remaining_time}").font.bold = False  # Make font bold
                                    bold_paragraph.runs[-1].font.size = Pt(10)  # Set font size for the remaining times
                                    bold_paragraph.paragraph_format.line_spacing_rule = 0  # Set line spacing to
                                    bold_paragraph.paragraph_format.space_after = Pt(1)  # Reduce spacing after the paragraph
        

                    target_row[2].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[2].paragraphs[0].runs[0].font.size = Pt(10)
                    target_row[2].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[2].paragraphs[0].paragraph_format.space_after = Pt(1)  ### Reduce spacing after the paragraph

                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[2].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months

                    elif int(month) == 4 and int(date) < end_dst_day:
                        # print(f"Date: {date}, Month: {month}, Day: {day}", {end_dst_day})  # Print the year, month, and day for debugging
                        target_row[2].paragraphs[0].runs[0].font.bold = True  # Make font bold for April
                        target_row[2].paragraphs[0].paragraph_format.space_after = Pt(1) ###

                    elif int(month) == 4 and int(date) == end_dst_day:
                        target_row[2].paragraphs[0].runs[0].font.bold = False
                        bold_paragraph = target_row[2].paragraphs[0]
                        # Calculate rows where time is before 3 AM
                        times = [rest[i] for i in range(0, len(rest), 2)]
                        before_3am_count = sum(1 for time in times if time and int(time.split(':')[0]) < 3)
                        # Update the text for the cell based on the row count
                        ontimes = [rest[i] for i in range(1, len(rest), 2)][:before_3am_count] #range(0 to 1 for update the tide and direction)
                        remaining_times = [rest[i] for i in range(1, len(rest), 2)][before_3am_count:]

                        if ontimes:

                            bold_paragraph.clear()  # Clear existing text
                            bold_paragraph.runs.clear()  # Clear existing runs
                            run = bold_paragraph.add_run("\n".join(ontimes))
                            run.font.size = Pt(10)

                            if no_condition == True:
                                run.font.bold = False
                            else:
                                run.font.bold = True

                            for remaining_time in remaining_times:
                                if remaining_time:
                                    bold_paragraph.add_run(f"\n{remaining_time}").font.bold = False  # Make font bold
                                    bold_paragraph.runs[-1].font.size = Pt(10)  # Set font size for the remaining times
                                    bold_paragraph.paragraph_format.line_spacing_rule = 0  # Set line spacing to
                                    bold_paragraph.paragraph_format.space_after = Pt(1)  # Reduce spacing after the paragraph
                       
        

            for idx, row_date in enumerate(range(9, 17)):  # Changed range to 9 to 16 for the next 8 days
                if date == str(row_date):
                    # Fill the second column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the second column
                    target_row[3].text = f"{date}"
                    run_date = target_row[3].paragraphs[0].runs[0]
                    run_date.font.size = Pt(22)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[3].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(11)
                    target_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[3].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[3].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    target_row[4].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[4].paragraphs[0].runs[0].font.size = Pt(10)
                    target_row[4].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[4].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[4].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                    
                    target_row[5].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[5].paragraphs[0].runs[0].font.size = Pt(10)
                    target_row[5].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[5].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[5].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months

            for idx, row_date in enumerate(range(17, 25)):  # Changed range to 17 to 24 for the next 8 days
                if date == str(row_date):
                    # Fill the third column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the third column
                    target_row[6].text = f"{date}"
                    run_date = target_row[6].paragraphs[0].runs[0]
                    run_date.font.size = Pt(22)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[6].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(11)
                    target_row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[6].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[6].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    target_row[7].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[7].paragraphs[0].runs[0].font.size = Pt(10)
                    target_row[7].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[7].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[7].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                    
                    target_row[8].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[8].paragraphs[0].runs[0].font.size = Pt(10)
                    target_row[8].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[8].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[8].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months

            for idx, row_date in enumerate(range(25, 32)):  # Changed range to 25 to 31 for the next 8 days
                if date == str(row_date):
                    # Fill the fourth column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the fourth column
                    target_row[9].text = f"{date}"
                    run_date = target_row[9].paragraphs[0].runs[0]
                    run_date.font.size = Pt(22)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[9].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(11)
                    target_row[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[9].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[9].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

                    target_row[10].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[10].paragraphs[0].runs[0].font.size = Pt(10)
                    target_row[10].paragraphs[0].paragraph_format.line_spacing_rule = 0  # Set line spacing to single
                    target_row[10].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
                    
                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[10].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                    elif int(month) == 9 and int(date) > start_dst_day:
                        # print(f"Date: {date}, Month: {month}, Day: {day}", {start_dst_day})
                        target_row[10].paragraphs[0].runs[0].font.bold = True  # Make font bold for September
                    elif int(month) == 9 and int(date) == start_dst_day:
        
                        target_row[10].paragraphs[0].runs[0].font.bold = True  # Make font bold for September on the start of DST
        
                        # Calculate rows where time is after 2 AM
                        times = [rest[i] for i in range(0, len(rest), 2)]
                        before_2am_count = sum(1 for time in times if time and int(time.split(':')[0]) < 2)

                        # Update the text for the cell based on the DTS count
                        bold_paragraph = target_row[10].paragraphs[0]

                        ontimes = [rest[i] for i in range(0, len(rest), 2)][:before_2am_count]
                        remaining_times = [rest[i] for i in range(0, len(rest), 2)][before_2am_count:]

                        if ontimes:
                            bold_paragraph.clear()  # Clear existing text
                            bold_paragraph.runs.clear()  # Clear existing runs
                            run = bold_paragraph.add_run("\n".join(ontimes))
                            run.font.size = Pt(10)

                            for remaining_time in remaining_times:
                                if remaining_time:
                                    bold_paragraph.add_run(f"\n{remaining_time}").font.bold = True  # Make font bold
                                    bold_paragraph.runs[-1].font.size = Pt(10)  # Set font size for the remaining times
                                    bold_paragraph.paragraph_format.line_spacing_rule = 0  # Set line spacing to
                                    bold_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

                    target_row[11].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[11].paragraphs[0].runs[0].font.size = Pt(10)
                    target_row[11].paragraphs[0].paragraph_format.line_spacing_rule = 0
                    target_row[11].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

                    if int(month) in [1, 2, 3, 10, 11, 12]:
                        target_row[11].paragraphs[0].runs[0].font.bold = True  # Make font bold for specific months
                    elif int(month) == 9 and int(date) > start_dst_day:
                        # print(f"Date: {date}, Month: {month}, Day: {day}", {start_dst_day})
                        target_row[11].paragraphs[0].runs[0].font.bold = True  # Make font bold for September

                    elif int(month) == 9 and int(date) == start_dst_day:

                        target_row[11].paragraphs[0].runs[0].font.bold = True
                        bold_paragraph = target_row[11].paragraphs[0]
                        # Calculate rows where time is after 2 AM
                        times = [rest[i] for i in range(0, len(rest), 2)]
                        before_2am_count = sum(1 for time in times if time and int(time.split(':')[0]) < 2)
                        
                        ontimes = [rest[i] for i in range(1, len(rest), 2)][:before_2am_count] #range(0 to 1 for update the tide and direction)
                        remaining_times = [rest[i] for i in range(1, len(rest), 2)][before_2am_count:]

                        # print(f"ontimes: {ontimes}, remaining_times: {remaining_times}")  # Debugging output

                        if ontimes:
                            bold_paragraph.clear()  # Clear existing text
                            bold_paragraph.runs.clear()  # Clear existing runs
                            run = bold_paragraph.add_run("\n".join(ontimes))
                            run.font.size = Pt(10)

                            for remaining_time in remaining_times:
                                if remaining_time:
                                    bold_paragraph.add_run(f"\n{remaining_time}").font.bold = True  # Make font bold
                                    bold_paragraph.runs[-1].font.size = Pt(10)  # Set font size for the remaining times
                                    bold_paragraph.paragraph_format.line_spacing_rule = 0  # Set line spacing to
                                    bold_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph

            # # Set row height for table rows starting from row number 2 (index 1)
            for tbl_row in table.rows[1:]:
                tbl_row.height = Pt(55)
                tbl_row.height_rule = 1  # 1 = At least, 2 = Exactly (optional, can use 1 for "at least")

        # Add spacing after the table
        document.add_paragraph().paragraph_format.space_after = Pt(12)

        if region_name == "Te Aumiti / French Pass" or region_name == "Tory Channel / Kura Te Au Entrance":
            add_caution(document)

        if region_name == "Owenga - Chatham Island" or region_name == "Kaingaroa - Chatham Island" or region_name == "Waitangi - Chatham Island":
            add_daylight1(document, month)  # Add Chatham Islands daylight line after the table
        else:
            add_daylight(document, month)  # Add NZ daylight line after the table

        add_copyright(document)  # Add copyright line after the table
    
    # Check if the file exists and remove it
    if os.path.exists(output_path):
        os.remove(output_path)
    
    # Save the document
    document.save(output_path)

def convert_to_pdf(docx_path, pdf_path):
    """
    Convert a Word document to a PDF file.
    """
    try:
        convert(docx_path, pdf_path)
    except FileNotFoundError:
        print(f"Error: The file '{docx_path}' does not exist.")
    except Exception as e:
        print(f"An error occurred while converting to PDF: {e}")

def load_config():
    """
    Load configuration from the 'config.yaml' file.

    Returns:
        tuple: A tuple containing folder_path and output_folder.
    """
    with open('config.yaml', 'r') as config_file:
        config = yaml.safe_load(config_file)

    folder_path = config['folder_path']
    output_folder = config['output_folder']
    linz_logo_path = config['linz_logo_path']
    os.makedirs(output_folder, exist_ok=True)

    return folder_path, output_folder, linz_logo_path

def main():
    """Main function to execute the script."""
    # Define the folder path containing CSV files
    # Load configuration from config.yaml
    folder_path, output_folder, linz_logo_path = load_config()
    # # Check if the folder exists
    # if not os.path.exists(folder_path):
    #     print(f"Error: The folder '{folder_path}' does not exist.")
    #     return
    # # Check if the output folder exists, if not create it
    # if not os.path.exists(output_folder):
    #     os.makedirs(output_folder)
    # # Check if the output folder is empty
    # if os.listdir(output_folder):
    #     print(f"Error: The output folder '{output_folder}' is not empty.")
    #     return

    # Process each CSV file in the folder
    for file in os.listdir(folder_path):
        if file.endswith('.csv'):
            file_path = os.path.join(folder_path, file)
            output_path = os.path.join(output_folder, os.path.splitext(file)[0] + '.docx')
            pdf_path = os.path.join(output_folder, os.path.splitext(file)[0] + '.pdf')

            try:
                # Read the CSV file
                file_info, header, data = read_csv(file_path)

                # Group data by month
                grouped_data = group_data_by_month(data)

                # Save grouped data to a Word document
                save_to_word(file_info, grouped_data, output_path, linz_logo_path)

                # Convert the Word document to PDF
                convert_to_pdf(output_path, pdf_path)

                print(f"Processed: {file}")
                print(f"Word document saved to {output_path}")
                print(f"PDF document saved to {pdf_path}")
            except FileNotFoundError:
                print(f"Error: The file '{file_path}' does not exist.")
            except ValueError as ve:
                print(f"ValueError while processing '{file}': {ve}")
            except Exception as e:
                print(f"An unexpected error occurred while processing '{file}': {e}")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"An error occurred in the main function: {e}")