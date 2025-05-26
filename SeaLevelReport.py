import csv
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_csv(file_path):
    data = []
    with open(file_path, mode='r') as file:
        csv_reader = csv.reader(file)
        header = next(csv_reader)
        for row in csv_reader:
            data.append(row)
    return header, data

# Example usage
# file_path = r'C:\Projects\Jenn\2025\Te Aumiti - French Pass_2024.csv'
file_path = r'C:\Projects\Glen\CSV files\Bluff_2022-23_NZNA_DT.csv'
header, data = read_csv(file_path)
regionname = header[1]
coord = (f"Lat. {header[2]} Long. {header[3]}")

print(coord)

monthly_data = defaultdict(list)
for row in data[2:]:
    # Replace empty values with ''
    row = [value if value else '' for value in row]
    
    if len(row) == 10:
        date, day, month, year, t1, d1, t2, d2, t3, d3 = row
        # print(f"Data for month {month}:")
    elif len(row) == 12:
        date, day, month, year, t1, d1, t2, d2, t3, d3, t4, d4 = row
        # print(f"Data for month {month}:")
    elif len(row) == 13:
        date, day, month, year, t1, d1, t2, d2, t3, d3, t4, d4, t5, d5 = row
        # print(f"Data for month {month}:")
    else:
        print("Row does not have exactly 10 or 12 or 14 elements:", row)
    monthly_data[month].append(row)



def save_to_word(data, file_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'  # Set font to Arial

    # Function to add the top table
    def add_top_table(doc):
        top_table = doc.add_table(rows=1, cols=2)
        # top_table.style = 'Table Grid'

        # Add logo to the first cell
        top_table.cell(0, 0).paragraphs[0].add_run().add_picture(r'N:\Publications\Toitū Te Whenua LINZ logo\toitu_te_whenua_colour_cmyk_66mm_png.png', width=Pt(180))
        top_table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(0)  # Remove blank space after paragraph


        # Add text and hyperlinks to the second cell
        text_paragraph = top_table.cell(0, 1).paragraphs[0]
        text_content = [
            ("Sourced from ", None),
            ("http://www.linz.govt.nz", RGBColor(0, 0, 255)),
            ("\nE-mail address ", None),
            ("hydro@linz.govt.nz", RGBColor(0, 0, 255))
        ]
        for text, color in text_content:
            run = text_paragraph.add_run(text)
            run.font.size = Pt(12)
            if color:
                run.font.color.rgb = color
            run.font.underline = True
        text_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the paragraph
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align text to the left

        # Format cells
        for cell in top_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add the top table to the first page
    add_top_table(doc)

    # Add some spacing after the top table
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)


    # Function to add a title line to the document
    def add_title(doc):
        title_paragraph = doc.add_paragraph("New Zealand Hydrographic Authority Tide Stream Predictions")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.name = 'Arial'
        title_paragraph.paragraph_format.space_after = Pt(2)  # Reduce spacing after the title

    # Add the title line to the first page
    add_title(doc)

    # Add a regionname to the document
    def add_header(doc, regionname):
        header_paragraph = doc.add_paragraph(regionname)
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_paragraph.runs[0]
        header_run.font.size = Pt(20)
        header_run.font.bold = True
        header_run.font.name = 'Arial'
        header_run.font.color.rgb = RGBColor(20, 171, 155)  # Set color to #14ab9b
        header_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the header

    # Add coordinates to the document
    add_header(doc, regionname)

    def add_coordinates(doc, coord):
        ldt_paragraph = doc.add_paragraph(coord)
        ldt_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ldt_run = ldt_paragraph.runs[0]
        ldt_run.font.size = Pt(11)
        ldt_run.font.name = 'Arial'
        ldt_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the header

    # Add coordinates to the document
    add_coordinates(doc, coord)

    def add_date(doc, month_data):
        month_names = ["January", "February", "March", "April", "May", "June", 
                   "July", "August", "September", "October", "November", "December"]
        month_name = month_names[int(month_data[1][2]) - 1]
        date_paragraph = doc.add_paragraph(f"{month_name} {month_data[1][3]}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.runs[0]
        date_run.font.size = Pt(20)
        date_run.font.bold = True
        date_run.font.name = 'Arial'
        date_run.font.color.rgb = RGBColor(20, 171, 155)  # Set color to #14ab9b
        date_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the header

    # Add date to the document
    add_date(doc, monthly_data['1'])
    
    # Function to add a condition line to the document
    def add_condition(doc):
        condition_paragraph = doc.add_paragraph("Tidal Stream begins at the N.Z. Local Time shown, in the direction indicated")
        condition_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        condition_run = condition_paragraph.runs[0]
        condition_run.font.size = Pt(11)
        condition_run.font.name = 'Arial'
        condition_paragraph.paragraph_format.space_after = Pt(5)  # Reduce spacing after the title

    # Add the title line to the first page
    add_condition(doc)


    # Add data table
    table = doc.add_table(rows=9, cols=12)
    hdr_cells = table.rows[0].cells
    headers = ['', 'Time', 'Dir', '', 'Time', 'Dir', '', 'Time', 'Dir', '', 'Time', 'Dir']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing under the header
            for run in paragraph.runs:
                run.font.size = Pt(10)  # Set header font size to 10
    
    for month in monthly_data.keys():
        month_data = monthly_data[month]
        for row in month_data:
            date, day, month, year, *rest = row
            for idx, row_date in enumerate(range(1, 9)):  # Changed range to 1 to 9 for 8 days
                if date == str(row_date):
                    # Fill the first column with the first 8 days
                    target_row = table.rows[idx + 1].cells
                    target_row[0].text = f"{date}"
                    run_date = target_row[0].paragraphs[0].runs[0]
                    run_date.font.size = Pt(20)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[0].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(10)
                    target_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[0].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing

                    target_row[1].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[1].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[1].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing
                    target_row[2].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[2].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[2].paragraphs[0].paragraph_format.space_after = Pt(3)  # Reduce spacing

            for idx, row_date in enumerate(range(9, 17)):  # Changed range to 9 to 16 for the next 8 days
                if date == str(row_date):
                    # Fill the second column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the second column
                    target_row[3].text = f"{date}"
                    run_date = target_row[3].paragraphs[0].runs[0]
                    run_date.font.size = Pt(20)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[3].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(10)
                    target_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[3].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing

                    target_row[4].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[4].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[4].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing
                    target_row[5].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[5].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[5].paragraphs[0].paragraph_format.space_after = Pt(3)  # Reduce spacing

            for idx, row_date in enumerate(range(17, 25)):  # Changed range to 17 to 24 for the next 8 days
                if date == str(row_date):
                    # Fill the third column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the third column
                    target_row[6].text = f"{date}"
                    run_date = target_row[6].paragraphs[0].runs[0]
                    run_date.font.size = Pt(20)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[6].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(10)
                    target_row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[6].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing

                    target_row[7].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[7].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[7].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing
                    target_row[8].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[8].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[8].paragraphs[0].paragraph_format.space_after = Pt(3)  # Reduce spacing

            for idx, row_date in enumerate(range(25, 32)):  # Changed range to 25 to 31 for the next 8 days
                if date == str(row_date):
                    # Fill the fourth column with the next 8 days
                    target_row = table.rows[idx + 1].cells  # Adjust row index for the fourth column
                    target_row[9].text = f"{date}"
                    run_date = target_row[9].paragraphs[0].runs[0]
                    run_date.font.size = Pt(20)
                    run_date.font.bold = True  # Make the date font bold
                    run_day = target_row[9].paragraphs[0].add_run(f"\n{day}")
                    run_day.font.size = Pt(10)
                    target_row[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align center
                    target_row[9].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing

                    target_row[10].text = "\n".join([rest[i] for i in range(0, len(rest), 2)])
                    target_row[10].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[10].paragraphs[0].paragraph_format.space_after = Pt(0)  # Reduce spacing
                    target_row[11].text = "\n".join([rest[i] for i in range(1, len(rest), 2)])
                    target_row[11].paragraphs[0].runs[0].font.size = Pt(9)
                    target_row[11].paragraphs[0].paragraph_format.space_after = Pt(3)  # Reduce spacing

        # Function to add a caution line to the document
        def add_caution(doc):
            caution_paragraph = doc.add_paragraph("Caution: Tidal Streams may be subject to irregularities and these times should be regarded as approximate only.")
            caution_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caution_run = caution_paragraph.runs[0]
            caution_run.font.size = Pt(8.5)
            caution_run.font.name = 'Arial'
            caution_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the title

         # Add the title line to the first page
        add_caution(doc)

                # Function to add a daylight line to the document
        def add_daylight(doc):
            daylight_paragraph = doc.add_paragraph("Times listed are N.Z. Daylight Time")
            daylight_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            daylight_run = daylight_paragraph.runs[0]
            daylight_run.font.size = Pt(10)
            daylight_run.font.name = 'Arial'
            daylight_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the title

         # Add the title line to the first page
        add_daylight(doc)

        def add_copyright(doc):
            copyright_paragraph = doc.add_paragraph("Crown Copyright Reserved")
            copyright_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            copyright_run = copyright_paragraph.runs[0]
            copyright_run.font.size = Pt(10)
            copyright_run.font.name = 'Arial'
            copyright_paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing after the title

         # Add the title line to the first page
        add_copyright(doc)

        # Add a page break for the next month, continue until month is 12
        if int(month) < 12:
            doc.add_page_break()
            add_top_table(doc)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_title(doc)
            add_header(doc, regionname)
            add_coordinates(doc, coord)
            add_date(doc, monthly_data[str(int(month) + 1)])
            add_condition(doc)
            table = doc.add_table(rows=9, cols=12)
            hdr_cells = table.rows[0].cells
            headers = ['', 'Time', 'Dir', '', 'Time', 'Dir', '', 'Time', 'Dir', '', 'Time', 'Dir']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)  # Reduce spacing under the header
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # Set header font size to 10

    doc.save(file_path)

# Save the data to a Word document
output_file_path = r'C:\Projects\Jenn\2025\SeaLevelReport.docx'
save_to_word([header] + data, output_file_path)